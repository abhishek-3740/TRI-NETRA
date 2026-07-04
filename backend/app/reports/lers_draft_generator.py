"""
LERS (Legal Emergency Request System) DRAFT generator.

IMPORTANT: this produces a draft document only. It never sends anything
over a network — an officer must review, sign, and submit it through
proper legal channels. Auto-dispatching legal requests would both break
the air-gapped design and create an unreviewed legal-instrument liability.
"""
from pathlib import Path
from docx import Document

REPORT_DIR = Path("ml_artifacts") / "reports"


def generate_lers_draft(case_id: str) -> str:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = REPORT_DIR / f"LERS_DRAFT_{case_id}.docx"

    # TODO: pull real suspect phone/IMEI/date-range from Postgres for this case
    case_data = _get_case_data(case_id)

    doc = Document()
    doc.add_heading("LERS Request — DRAFT (Requires Officer Review & Signature)", level=1)
    doc.add_paragraph(f"Case ID: {case_id}")
    doc.add_paragraph(f"Suspect Phone Number: {case_data.get('phone_number', '[TO FILL]')}")
    doc.add_paragraph(f"IMEI: {case_data.get('imei', '[TO FILL]')}")
    doc.add_paragraph(f"Requested Date Range: {case_data.get('date_range', '[TO FILL]')}")
    doc.add_paragraph(f"Requesting Officer: {case_data.get('officer_name', '[TO FILL]')}")
    doc.add_paragraph("")
    doc.add_paragraph(
        "NOTE: This is an auto-generated DRAFT only. It must be reviewed, "
        "signed, and submitted manually through the appropriate legal channel. "
        "This system does not transmit this request automatically."
    )

    doc.save(output_path)
    return str(output_path)


def _get_case_data(case_id: str) -> dict:
    return {"case_id": case_id}
